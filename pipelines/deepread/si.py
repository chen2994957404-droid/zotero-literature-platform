# -*- coding: utf-8 -*-
"""SI（补充材料）精读这一步：SI 附件 → 「实验细节精读」HTML。

原来是 `文献精读/si_deepread.py` 的 `main()`，被 watcher 用 subprocess 拉起来。
搬进来之后是一个能直接调、能直接测的函数，行为不变。

定位（与正文精读的分工，用户 2026-07-25 定）：
    正文精读 = 理解这篇做了什么；SI 精读 = 我要复现时查参数。
SI 里有正文完全没有的可复现细节：精确投料量、原料分子量、溶剂配比、对照组设计逻辑。
"""
import base64
import io
import os
import re
import zipfile

from adapters.llm_client import chat
from adapters.pdf_parse import parse_pdf, PDFParseError
from adapters.zotero_client import zget, USER_ID, STORAGE_DIR, SUPP_PAT
from core import paths
from core.config import get_key
from domain.figure_crop import crop_figures
from domain.si_filter import filtered_text

PROMPT_VER = 1
PRODUCER = 'si_deepread'

DOCX_CT = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

SYS = """你是材料科学文献助手。下面是一篇论文【补充材料(SI)】的正文。
请生成中文的"实验细节精读"，服务于"我要复现这个实验"的读者。要求：
1. 【原料与规格】所有试剂及关键规格（尤其分子量、纯度、供应商）
2. 【合成步骤】逐步写清：投料量(克数/摩尔数)、配比、溶剂、温度、时间、后处理。最重要，务必精确
3. 【对照组设计】多个样品时说明设计逻辑与差异
4. 【表征方法】一句话列出手段即可，不展开仪器型号
5. 【补充数据要点】SI 图表说明的关键结论，按【图N】顺序简述
讨论到某张补充图时用【图N】标记插图位置。用中文，专业准确，数值保留原文单位。"""


class SIFailed(Exception):
    """SI 这一步没能产出。与「这篇根本没有 SI」是两回事，别混。"""


def find_si_file(item_key):
    """定位 SI 附件文件。支持 PDF 和 .docx（Elsevier 的 SI 常是 docx）。

    返回 (路径, 类型)，类型 ∈ 'pdf'/'docx'；找不到返回 (None, None)。
    """
    try:
        children = zget(f'/users/{USER_ID}/items/{item_key}/children')
    except Exception:
        return None, None
    for c in children:
        d = c['data']
        if d.get('itemType') != 'attachment':
            continue
        if d.get('contentType', '') not in ('application/pdf', DOCX_CT):
            continue
        title = (d.get('title') or '').strip()
        fn = (d.get('filename') or '')
        if not (SUPP_PAT.search(title) or SUPP_PAT.search(fn) or title.upper() == 'SI'):
            continue
        dd = os.path.join(STORAGE_DIR, c['key'])
        if os.path.isdir(dd):
            for f in os.listdir(dd):
                if f.lower().endswith('.pdf'):
                    return os.path.join(dd, f), 'pdf'
                if f.lower().endswith('.docx'):
                    return os.path.join(dd, f), 'docx'
    return None, None


def read_docx_text(path):
    """读 .docx 的文字（含表格）。表格常含关键参数，务必取。"""
    try:
        import docx
    except ImportError:
        raise SIFailed('需要 python-docx：pip install python-docx')
    doc = docx.Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tb in doc.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(' | '.join(cells))
    return '\n\n'.join(parts)


def extract_docx_images(path, min_kb=15, log=print):
    """取出 .docx 里内嵌的图片（docx 本质是 zip，图在 word/media/）。

    返回 [{b64, caption, num}]，格式与 figure_crop 一致，可直接进渲染流程。
    过滤掉小图标（<min_kb），只留有意义的补充图。
    """
    figs = []
    try:
        with zipfile.ZipFile(path) as z:
            media = sorted(n for n in z.namelist()
                           if n.startswith('word/media/')
                           and n.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')))
            for name in media:
                data = z.read(name)
                if len(data) < min_kb * 1024:      # 滤掉小图标/装饰
                    continue
                ext = name.rsplit('.', 1)[-1].lower()
                mime = 'jpeg' if ext in ('jpg', 'jpeg') else ext
                figs.append({
                    'b64': f'data:image/{mime};base64,' + base64.b64encode(data).decode(),
                    'caption': '', 'num': len(figs) + 1})
    except Exception as e:
        log(f'  （docx 取图失败，跳过：{e}）')
    return figs


CSS = ('body{max-width:820px;margin:0 auto;padding:24px;font-family:-apple-system,'
       '"Microsoft YaHei",sans-serif;line-height:1.85;color:#222;background:#fafafa}'
       'h2.section{background:linear-gradient(90deg,#e8934a,#d4703a);color:#fff;padding:8px 20px;'
       'border-radius:20px;display:inline-block;font-size:19px;margin:34px 0 16px}'
       'h3{color:#c26a35;font-size:16px;margin-top:22px}p{margin:12px 0;text-align:justify}'
       'img{max-width:100%;display:block;margin:18px auto;border:1px solid #eee;'
       'border-radius:6px;box-shadow:0 2px 8px rgba(0,0,0,.06)}strong{color:#c0392b}')


def render_html(content, figs, title=''):
    """把 Markdown 式内容 + 图渲染成 HTML（复用精读线的确定性插图思路）。"""
    used = set()

    def repl(m):
        n = int(m.group(1))
        used.add(n)
        return f'\n<img src="{figs[n-1]["b64"]}">\n' if 1 <= n <= len(figs) else ''

    content = re.sub(r'【图(\d+)】', repl, content)
    missing = [i for i in range(1, len(figs) + 1) if i not in used]
    if missing:
        content += '\n\n（其余补充图）\n' + ''.join(
            f'\n<img src="{figs[i-1]["b64"]}">\n' for i in missing)
    out = []
    for ln in content.split('\n'):
        s = ln.strip()
        if s.startswith('<img'):
            out.append(s)
            continue
        s = re.sub(r'^#{4,6}\s*', '', s)
        if s.startswith('### '):
            out.append(f'<h3>{s[4:].strip()}</h3>')
            continue
        if s.startswith('## '):
            out.append(f'<h2 class="section">{s[3:].strip()}</h2>')
            continue
        if s.startswith('# '):
            out.append(f'<h2 class="section">{s[2:].strip()}</h2>')
            continue
        s = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', s)
        s = re.sub(r'^[-*]\s+', '· ', s)
        if s:
            out.append(f'<p>{s}</p>')
    return ('<!DOCTYPE html><html lang="zh"><head><meta charset="utf-8">'
            f'<title>SI实验细节精读</title><style>{CSS}</style></head><body>'
            '<h2 class="section">补充材料（SI）· 实验细节精读</h2>'
            + '\n'.join(out) + '</body></html>')


MIN_OK = 800        # SI 精读低于这个字数基本是废品（正文线同款底线）

# 额度阶梯：**V4 的推理链计入 max_tokens**，给少了会「输出被截断且正文近乎为空」。
# 正文线早就是 32000 起 + 重试（血泪教训第 1 条），SI 线却一直是 6000 ——
# 于是同一篇有时成、有时败，看起来像玄学。2026-08-27 真实撞上一次才发现。
_BUDGETS = (16000, 32000)


def _call_llm(user, model, log=print):
    """调模型，额度不够就加倍再来一次。两次都不行才算失败，不写废品上盘。"""
    last = ''
    for i, budget in enumerate(_BUDGETS, 1):
        try:
            out = chat(SYS, user, provider='deepseek', model=model,
                       key=get_key('DEEPSEEK_KEY'), temperature=0.3,
                       max_tokens=budget)
        except Exception as e:
            log(f'  第{i}次调用失败（额度 {budget}）：{str(e)[:120]}')
            continue
        if len(out.strip()) >= MIN_OK:
            return out
        last = out
        log(f'  第{i}次输出仅 {len(out.strip())} 字（<{MIN_OK}），加大额度重试…')
    if len(last.strip()) >= MIN_OK:
        return last
    raise SIFailed(f'SI 精读输出仅 {len(last.strip())} 字，判定失败，不写盘')


def read_si(key, out_html=None, model=None, log=print):
    """跑一篇的 SI 精读。

    返回产物路径；**这篇根本没有 SI 附件时返回 None**（不是失败，别当失败记）。
    真出了问题抛 `SIFailed`。
    """
    key = paths.check_key(key)
    out_html = out_html or paths.si_summary(key)
    model = model or os.environ.get('SI_MODEL', 'deepseek-v4-flash')  # 输出长 → flash 省钱

    si_file, kind = find_si_file(key)
    if not si_file:
        log(f'[跳过] {key} 没有 SI 附件')
        return None
    log(f'[SI] {os.path.basename(si_file)} ({kind})')

    parsed = paths.si_parsed_dir(key)
    figs = []
    if kind == 'pdf':
        try:
            parse_pdf(si_file, parsed)      # 已解析则复用
        except PDFParseError as e:
            raise SIFailed(f'SI 解析失败：{e}')
        md = os.path.join(parsed, 'full.md')
        if not os.path.exists(md):
            raise SIFailed('SI 解析未生成 full.md')
        raw = io.open(md, encoding='utf-8').read()
        figs = crop_figures(parsed)
    else:                                   # docx：读文字（含表格）+ 取内嵌图片
        raw = read_docx_text(si_file)
        os.makedirs(parsed, exist_ok=True)
        io.open(os.path.join(parsed, 'full.md'), 'w', encoding='utf-8').write(raw)
        figs = extract_docx_images(si_file, log=log)
        log(f'  docx 读出 {len(raw)} 字符（含表格），取出内嵌图 {len(figs)} 张')

    body = filtered_text(raw)          # 过滤噪声（作者/单位/目录/参考文献）
    log(f'  过滤后 {len(body)} 字符（原 {len(raw)}），补充图 {len(figs)} 张')

    user = f'补充材料共有 {len(figs)} 张图。\n\n正文:\n{body[:30000]}'
    content = _call_llm(user, model, log)
    os.makedirs(os.path.dirname(out_html), exist_ok=True)
    io.open(out_html, 'w', encoding='utf-8').write(render_html(content, figs))
    log(f'  [完成] {out_html}  {round(os.path.getsize(out_html)/1024)} KB')
    return out_html
